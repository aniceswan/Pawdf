"""Shared pytest fixtures: generate small sample PDFs/DOCX on the fly so the
repo doesn't need to carry binary fixtures for the common case.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def _write_pdf(path: Path, num_pages: int, label_prefix: str = "Page") -> Path:
    c = canvas.Canvas(str(path), pagesize=letter)
    for i in range(1, num_pages + 1):
        c.drawString(72, 720, f"{label_prefix} {i}")
        c.showPage()
    c.save()
    return path


@pytest.fixture
def make_pdf(tmp_path: Path):
    """Factory fixture: make_pdf(name, num_pages) -> Path to a generated PDF."""

    def _make(name: str, num_pages: int, label_prefix: str = "Page") -> Path:
        return _write_pdf(tmp_path / name, num_pages, label_prefix)

    return _make


@pytest.fixture
def sample_pdf(make_pdf) -> Path:
    return make_pdf("sample.pdf", 5)


@pytest.fixture
def sample_pdf_with_image(tmp_path: Path) -> Path:
    from PIL import Image

    # Noisy (photo-like) image, not a solid color: solid colors compress
    # better losslessly than as JPEG, which would defeat the point of this
    # fixture (testing that lossy recompression shrinks a realistic image).
    noise = Image.effect_noise((2000, 1500), 40)
    img_path = tmp_path / "embedded.png"
    noise.convert("RGB").save(img_path)

    pdf_path = tmp_path / "with_image.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawImage(str(img_path), 50, 400, width=400, height=300)
    c.showPage()
    c.save()
    return pdf_path


@pytest.fixture
def sample_images(tmp_path: Path) -> list[Path]:
    from PIL import Image

    paths = []
    for i, color in enumerate([(255, 0, 0), (0, 255, 0), (0, 0, 255)]):
        path = tmp_path / f"img{i}.png"
        Image.new("RGB", (64, 48), color=color).save(path)
        paths.append(path)
    return paths


@pytest.fixture
def structured_pdf(tmp_path: Path) -> Path:
    """A PDF with a real visual hierarchy: two heading sizes, body text that
    wraps across lines, and italic/bold runs.

    PDF -> Word has to infer all of that from font metrics and geometry, so a
    flat one-size document (like sample_pdf) can't exercise it.
    """
    import textwrap

    from reportlab.lib.pagesizes import A4

    path = tmp_path / "structured.pdf"
    width, height = A4
    c = canvas.Canvas(str(path), pagesize=A4)

    c.setFont("Helvetica-Bold", 26)
    c.drawString(60, height - 90, "Quarterly Report")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(60, height - 140, "1. Introduction")

    c.setFont("Helvetica", 11)
    body = (
        "This paragraph is long enough that it wraps across more than a single "
        "line in the rendered PDF, which lets us check that wrapped lines get "
        "merged back into one paragraph instead of many."
    )
    y = height - 170
    for line in textwrap.wrap(body, 78):
        c.drawString(60, y, line)
        y -= 15

    y -= 30
    c.setFont("Helvetica-Oblique", 11)
    c.drawString(60, y, "This line is italic.")
    y -= 30
    c.setFont("Helvetica-Bold", 11)
    c.drawString(60, y, "This line is bold.")

    c.showPage()
    c.setFont("Helvetica", 11)
    c.drawString(60, height - 100, "Content on the second page.")
    c.showPage()
    c.save()
    return path


@pytest.fixture
def encrypted_pdf(sample_pdf: Path, tmp_path: Path) -> Path:
    """A password-protected copy of sample_pdf.

    Worth generating rather than skipping: core/ tells "password-protected"
    apart from "corrupted" so the GUI can explain which one happened, and for
    the PDFium-backed paths that distinction comes from matching on an error
    message, which a library upgrade could silently reword.
    """
    import pikepdf

    path = tmp_path / "encrypted.pdf"
    with pikepdf.open(sample_pdf) as pdf:
        pdf.save(path, encryption=pikepdf.Encryption(user="secret", owner="secret"))
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Document", level=1)
    doc.add_paragraph("A plain paragraph for testing.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path
