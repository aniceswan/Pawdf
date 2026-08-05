from pawdf.core import docx_to_pdf
from pawdf.core._shared.pdf_io import page_count


def test_docx_to_pdf_produces_valid_pdf(sample_docx, tmp_path):
    out = docx_to_pdf.docx_to_pdf(sample_docx, tmp_path / "out.pdf")

    assert out.exists()
    assert page_count(out) >= 1


def test_docx_to_pdf_with_images_and_lists(tmp_path):
    from docx import Document
    from PIL import Image

    img_path = tmp_path / "pic.png"
    Image.new("RGB", (300, 200), color=(50, 100, 150)).save(img_path)

    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Some intro text.")
    doc.add_paragraph("Bullet one", style="List Bullet")
    doc.add_paragraph("Bullet two", style="List Bullet")
    doc.add_picture(str(img_path))
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "H1"
    table.cell(0, 1).text = "H2"
    table.cell(1, 0).text = "v1"
    table.cell(1, 1).text = "v2"

    docx_path = tmp_path / "report.docx"
    doc.save(str(docx_path))

    out = docx_to_pdf.docx_to_pdf(docx_path, tmp_path / "report.pdf")

    assert out.exists()
    assert out.stat().st_size > 0
    assert page_count(out) >= 1


def test_docx_to_pdf_with_page_break(tmp_path):
    from docx import Document
    from docx.enum.text import WD_BREAK

    doc = Document()
    doc.add_paragraph("Page one")
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph("Page two")

    docx_path = tmp_path / "organize.docx"
    doc.save(str(docx_path))

    out = docx_to_pdf.docx_to_pdf(docx_path, tmp_path / "organize.pdf")
    assert page_count(out) == 2
