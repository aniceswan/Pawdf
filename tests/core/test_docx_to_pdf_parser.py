from pawdf.core.docx_to_pdf.model import ImageIR, ListItemIR, ParagraphIR, TableIR
from pawdf.core.docx_to_pdf.parser import parse_docx


def test_parses_heading_and_paragraph_in_order(sample_docx):
    doc_ir = parse_docx(str(sample_docx))

    para_blocks = [b for b in doc_ir.blocks if isinstance(b, ParagraphIR)]
    assert any(b.style.startswith("Heading") for b in para_blocks)
    assert any("plain paragraph" in "".join(r.text for r in b.runs) for b in para_blocks)


def test_heading_appears_before_paragraph(sample_docx):
    doc_ir = parse_docx(str(sample_docx))
    para_blocks = [b for b in doc_ir.blocks if isinstance(b, ParagraphIR)]

    assert para_blocks[0].style.startswith("Heading")
    assert para_blocks[1].style == "Normal"


def test_parses_table_with_correct_dimensions(sample_docx):
    doc_ir = parse_docx(str(sample_docx))
    tables = [b for b in doc_ir.blocks if isinstance(b, TableIR)]

    assert len(tables) == 1
    table = tables[0]
    assert len(table.rows) == 2
    assert len(table.rows[0]) == 2


def test_table_cell_text_matches_source(sample_docx):
    doc_ir = parse_docx(str(sample_docx))
    table = next(b for b in doc_ir.blocks if isinstance(b, TableIR))

    first_cell_blocks = table.rows[0][0]
    first_para = next(b for b in first_cell_blocks if isinstance(b, ParagraphIR))
    assert "".join(r.text for r in first_para.runs) == "A1"


def test_run_formatting_is_captured(tmp_path):
    from docx import Document

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("bold text")
    run.bold = True
    path = tmp_path / "formatted.docx"
    doc.save(str(path))

    doc_ir = parse_docx(str(path))
    para = next(b for b in doc_ir.blocks if isinstance(b, ParagraphIR))
    assert para.runs[0].bold is True
    assert para.runs[0].text == "bold text"


def test_bulleted_list_is_detected(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("first item", style="List Bullet")
    doc.add_paragraph("second item", style="List Bullet")
    path = tmp_path / "list.docx"
    doc.save(str(path))

    doc_ir = parse_docx(str(path))
    list_items = [b for b in doc_ir.blocks if isinstance(b, ListItemIR)]

    assert len(list_items) == 2
    assert all(item.ordered is False for item in list_items)


def test_numbered_list_is_detected_as_ordered(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("first", style="List Number")
    path = tmp_path / "numbered.docx"
    doc.save(str(path))

    doc_ir = parse_docx(str(path))
    list_items = [b for b in doc_ir.blocks if isinstance(b, ListItemIR)]

    assert len(list_items) == 1
    assert list_items[0].ordered is True


def test_inline_image_is_extracted_with_bytes_and_size(tmp_path):
    from docx import Document
    from PIL import Image

    img_path = tmp_path / "pic.png"
    Image.new("RGB", (100, 50), color=(10, 20, 30)).save(img_path)

    doc = Document()
    doc.add_picture(str(img_path))
    path = tmp_path / "with_pic.docx"
    doc.save(str(path))

    doc_ir = parse_docx(str(path))
    images = [b for b in doc_ir.blocks if isinstance(b, ImageIR)]

    assert len(images) == 1
    assert len(images[0].data) > 0
    assert images[0].width_pt is not None
    assert images[0].height_pt is not None
