import pytest

from pawdf.core import pdf_to_docx


def _paragraphs(path):
    """Non-empty (style_name, text, any_bold, any_italic) tuples from a .docx."""
    from docx import Document

    doc = Document(str(path))
    return [
        (
            p.style.name,
            p.text,
            any(r.bold for r in p.runs),
            any(r.italic for r in p.runs),
        )
        for p in doc.paragraphs
        if p.text.strip()
    ]


def test_pdf_to_docx_produces_a_docx_file(sample_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(sample_pdf, tmp_path / "out.docx")

    assert out.exists()
    assert out.suffix == ".docx"
    assert out.stat().st_size > 0


def test_pdf_to_docx_produces_openable_document(sample_pdf, tmp_path):
    from docx import Document

    out = pdf_to_docx.pdf_to_docx(sample_pdf, tmp_path / "out.docx")

    doc = Document(str(out))
    assert doc is not None


def test_pdf_to_docx_supports_page_range(sample_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(sample_pdf, tmp_path / "out.docx", start_page=0, end_page=2)
    assert out.exists()


def test_page_range_limits_extracted_content(structured_pdf, tmp_path):
    full = _paragraphs(pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "full.docx"))
    first = _paragraphs(
        pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "first.docx", start_page=0, end_page=1)
    )

    assert any("second page" in text for _, text, _, _ in full)
    assert not any("second page" in text for _, text, _, _ in first)


def test_text_is_actually_extracted(structured_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx")
    all_text = " ".join(text for _, text, _, _ in _paragraphs(out))

    assert "Quarterly Report" in all_text
    assert "Introduction" in all_text


def test_larger_fonts_become_headings(structured_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx")
    styles = {text: style for style, text, _, _ in _paragraphs(out)}

    assert styles["Quarterly Report"].startswith("Heading")
    assert styles["1. Introduction"].startswith("Heading")
    # The 26pt title should outrank the 16pt section heading.
    assert styles["Quarterly Report"] < styles["1. Introduction"]


def test_body_text_is_not_promoted_to_a_heading(structured_pdf, tmp_path):
    paragraphs = _paragraphs(pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx"))

    body = [(style, text) for style, text, _, _ in paragraphs if "wraps across" in text]
    assert body, "expected the wrapped body paragraph to survive"
    assert body[0][0] == "Normal"


def test_wrapped_lines_merge_into_one_paragraph(structured_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx")

    wrapped = [text for _, text, _, _ in _paragraphs(out) if "wraps across" in text]
    assert len(wrapped) == 1
    # The whole sentence should land in that single paragraph, not be split.
    assert "instead of many." in wrapped[0]


def test_separate_lines_stay_separate_paragraphs(structured_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx")
    texts = [text for _, text, _, _ in _paragraphs(out)]

    assert "This line is italic." in texts
    assert "This line is bold." in texts


def test_bold_and_italic_are_preserved(structured_pdf, tmp_path):
    out = pdf_to_docx.pdf_to_docx(structured_pdf, tmp_path / "out.docx")
    by_text = {text: (bold, italic) for _, text, bold, italic in _paragraphs(out)}

    assert by_text["This line is italic."] == (False, True)
    assert by_text["This line is bold."] == (True, False)


def test_encrypted_pdf_raises_encrypted_error(encrypted_pdf, tmp_path):
    from pawdf.core._shared.errors import EncryptedPdfError

    with pytest.raises(EncryptedPdfError):
        pdf_to_docx.pdf_to_docx(encrypted_pdf, tmp_path / "out.docx")


def test_corrupt_pdf_raises_conversion_error(tmp_path):
    from pawdf.core._shared.errors import ConversionError

    bad = tmp_path / "bad.pdf"
    bad.write_bytes(b"definitely not a PDF")
    with pytest.raises(ConversionError):
        pdf_to_docx.pdf_to_docx(bad, tmp_path / "out.docx")
