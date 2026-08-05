"""DocumentIR -> .docx, via python-docx.

Consumes only the IR, so it can be tested by handing it a DocumentIR built
in code, with no PDF involved.
"""

from __future__ import annotations

import io
from pathlib import Path

from pawdf.core.pdf_to_docx.model import DocumentIR, ImageIR, ParagraphIR

# Word's Heading styles only go where we put them; anything deeper than this
# is clamped so an unusual document can't ask for a style that isn't defined.
_MAX_HEADING_LEVEL = 3


def _add_paragraph(document, paragraph: ParagraphIR) -> None:  # noqa: ANN001 - python-docx Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    if paragraph.heading_level > 0:
        level = min(paragraph.heading_level, _MAX_HEADING_LEVEL)
        block = document.add_heading("", level=level)
    else:
        block = document.add_paragraph()

    if paragraph.alignment == "center":
        block.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif paragraph.alignment == "right":
        block.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run_ir in paragraph.runs:
        run = block.add_run(run_ir.text)
        run.bold = run_ir.bold
        run.italic = run_ir.italic
        # Headings keep the size their Word style defines; overriding it with
        # the PDF's measured size would flatten the visual hierarchy that
        # heading detection just recovered.
        if paragraph.heading_level == 0 and run_ir.font_size_pt > 0:
            run.font.size = Pt(run_ir.font_size_pt)


def _add_image(document, image: ImageIR) -> None:  # noqa: ANN001 - python-docx Document
    from docx.shared import Pt

    try:
        document.add_picture(io.BytesIO(image.data), width=Pt(image.width_pt))
    except Exception:  # noqa: BLE001 - an image Word can't embed shouldn't fail the document
        pass


def render_to_docx(document_ir: DocumentIR, out_path: str | Path) -> Path:
    from docx import Document

    document = Document()
    for index, page in enumerate(document_ir.pages):
        if index > 0:
            document.add_page_break()
        for block in page.blocks:
            if isinstance(block, ParagraphIR):
                _add_paragraph(document, block)
            elif isinstance(block, ImageIR):
                _add_image(document, block)

    out_path = Path(out_path)
    document.save(str(out_path))
    return out_path
