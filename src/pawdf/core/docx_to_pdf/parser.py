"""python-docx Document -> DocumentIR.

Walks the document body in document order (paragraphs and tables
interleaved), not python-docx's separate `.paragraphs`/`.tables` lists,
which would lose the interleaving. See docs/conversion_fidelity.md for what
this parser does and doesn't attempt to capture.
"""

from __future__ import annotations

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from pawdf.core.docx_to_pdf.model import (
    BlockIR,
    DocumentIR,
    ImageIR,
    ListItemIR,
    PageBreakIR,
    ParagraphIR,
    RunIR,
    TableIR,
)

_ALIGNMENT_MAP = {
    0: "left",  # WD_ALIGN_PARAGRAPH.LEFT
    1: "center",
    2: "right",
    3: "justify",
}

_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_EMU_PER_POINT = 12700


def parse_docx(path: str) -> DocumentIR:
    document = Document(path)
    blocks = list(_parse_block_items(document, document))
    return DocumentIR(blocks=blocks)


def _iter_block_items(parent):
    """Yield Paragraph/Table children of `parent` (a Document or _Cell) in
    document order. Standard pattern from the python-docx FAQ, since
    `.paragraphs`/`.tables` alone don't preserve interleaving.
    """
    if hasattr(parent, "element"):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _parse_block_items(document: Document, parent) -> list[BlockIR]:
    blocks: list[BlockIR] = []
    for item in _iter_block_items(parent):
        if isinstance(item, Paragraph):
            blocks.extend(_parse_paragraph(document, item))
        elif isinstance(item, Table):
            blocks.append(_parse_table(document, item))
    return blocks


def _parse_paragraph(document: Document, paragraph: Paragraph) -> list[BlockIR]:
    blocks: list[BlockIR] = []

    for image in _extract_inline_images(document, paragraph):
        blocks.append(image)

    if _has_page_break(paragraph):
        blocks.append(PageBreakIR())

    para_ir = _paragraph_to_ir(paragraph)
    list_info = _list_info(paragraph)
    if list_info is not None:
        level, ordered = list_info
        blocks.append(ListItemIR(paragraph=para_ir, level=level, ordered=ordered))
    else:
        blocks.append(para_ir)

    return blocks


def _paragraph_to_ir(paragraph: Paragraph) -> ParagraphIR:
    runs = [_run_to_ir(r) for r in paragraph.runs if r.text]
    style_name = paragraph.style.name if paragraph.style is not None else "Normal"
    alignment = _ALIGNMENT_MAP.get(
        paragraph.alignment.value if paragraph.alignment is not None else 0, "left"
    )
    return ParagraphIR(runs=runs, style=style_name, alignment=alignment)


def _run_to_ir(run) -> RunIR:
    color_hex = None
    try:
        if run.font.color and run.font.color.rgb:
            color_hex = str(run.font.color.rgb)
    except AttributeError:
        pass

    font_size_pt = run.font.size.pt if run.font.size is not None else None

    return RunIR(
        text=run.text,
        bold=bool(run.bold),
        italic=bool(run.italic),
        underline=bool(run.underline),
        font_size_pt=font_size_pt,
        color_hex=color_hex,
    )


def _list_info(paragraph: Paragraph) -> tuple[int, bool] | None:
    """Returns (level, ordered) if this paragraph is a list item, else None.

    Two detection paths, since real-world .docx files vary in which one they
    use: explicit `numPr` on the paragraph (how Word itself writes lists),
    or a "List*" paragraph style with no numPr at all (how python-docx's own
    default template represents `style="List Bullet"`, verified empirically:
    it only sets `pStyle`, no numbering XML). `ordered` is a heuristic (style
    name contains "Number", or numPr's numFmt in the rare case that's
    reachable), not a full read of the numbering part. See
    conversion_fidelity.md.
    """
    p_pr = paragraph._p.pPr
    num_pr = p_pr.numPr if p_pr is not None else None
    style_name = (paragraph.style.name or "") if paragraph.style is not None else ""
    is_list_style = style_name.lower().startswith("list")

    if num_pr is None and not is_list_style:
        return None

    level = num_pr.ilvl.val if (num_pr is not None and num_pr.ilvl is not None) else 0
    ordered = "number" in style_name.lower()
    return level, ordered


def _has_page_break(paragraph: Paragraph) -> bool:
    for br in paragraph._p.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _extract_inline_images(document: Document, paragraph: Paragraph) -> list[ImageIR]:
    images: list[ImageIR] = []
    for blip in paragraph._p.iter(f"{{{_A_NS}}}blip"):
        r_id = blip.get(qn("r:embed"))
        if not r_id:
            continue
        try:
            part = document.part.related_parts[r_id]
            data = part.blob
        except KeyError:
            continue

        width_pt, height_pt = _find_extent_pt(blip)
        images.append(ImageIR(data=data, width_pt=width_pt, height_pt=height_pt))
    return images


def _find_extent_pt(blip) -> tuple[float | None, float | None]:
    """Walk up from a <a:blip> to the enclosing <wp:inline>/<wp:anchor> and
    read its <wp:extent> (EMUs), converting to points. Returns (None, None)
    if no extent is found (e.g. unusual/anchored drawing structures).
    """
    node = blip.getparent()
    while node is not None:
        tag = node.tag
        if tag in (f"{{{_WP_NS}}}inline", f"{{{_WP_NS}}}anchor"):
            extent = node.find(f"{{{_WP_NS}}}extent")
            if extent is not None:
                cx = extent.get("cx")
                cy = extent.get("cy")
                if cx and cy:
                    return int(cx) / _EMU_PER_POINT, int(cy) / _EMU_PER_POINT
            return None, None
        node = node.getparent()
    return None, None


def _cell_blocks(document: Document, cell: _Cell) -> list[BlockIR]:
    return _parse_block_items(document, cell)


def _parse_table(document: Document, table: Table) -> TableIR:
    rows: list[list[list[BlockIR]]] = []
    for row in table.rows:
        cells_blocks = [_cell_blocks(document, cell) for cell in row.cells]
        rows.append(cells_blocks)
    return TableIR(rows=rows)
